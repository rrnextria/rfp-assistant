from __future__ import annotations

import pytest


def test_parse_docx_returns_sections():
    """Test DOCX parser with a minimal DOCX created in memory."""
    import io
    from docx import Document

    # Create a test DOCX
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the introduction text.")
    doc.add_heading("Requirements", level=2)
    doc.add_paragraph("The vendor shall provide support.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    import sys
    sys.path.insert(0, "/home/ravi/git/rfp-assistant/services/content-service")
    from parser import parse_docx
    sections = parse_docx(buf.getvalue())

    assert len(sections) > 0
    texts = " ".join(s.text for s in sections)
    assert "introduction text" in texts.lower() or "vendor" in texts.lower()


def test_parse_docx_no_headings_returns_content():
    """Test DOCX parser falls back to paragraph text when no heading styles."""
    import io
    from docx import Document

    doc = Document()
    doc.add_paragraph("Just a paragraph with no heading.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    import sys
    sys.path.insert(0, "/home/ravi/git/rfp-assistant/services/content-service")
    from parser import parse_docx
    sections = parse_docx(buf.getvalue())
    assert len(sections) >= 1
    combined = " ".join(s.text for s in sections)
    assert "paragraph" in combined.lower()
